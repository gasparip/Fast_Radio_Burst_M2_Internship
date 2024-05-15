from docx import Document
from docx.shared import Inches,Cm,Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image
from matplotlib.pyplot import imshow
import sys
Listarg=[]
for arg in sys.argv:
    Listarg+= [arg]
#print(Listarg)

ForPDFtxt="/data/pgaspari/WeeklyReport/.processFile/ForPDFtxt.txt"
ForPythonToShToPDF="/data/pgaspari/WeeklyReport/.processFile/ForPythonToShToPDF.txt"
ToDoFile="/data/pgaspari/WeeklyReport/.processFile/ToDoFile.txt"
PathJpeg="/data/pgaspari/WeeklyReport/Result/"+Listarg[2]+"/JPEG/"
SavePath="/data/pgaspari/WeeklyReport/Result/"+Listarg[2]+"/DOCFile/Report_"+Listarg[2]+".odt"
# Créer un nouveau document
doc = Document()
# Lire le contenu du fichier texte
with open(ForPDFtxt, 'r') as file:
    contenu_texte = file.read()
with open(ForPythonToShToPDF, 'r') as file:
    contenu_texte2 = file.read() 
doc.add_heading("NenuFAR FRB Observation Weekly Report", level=1)
doc.add_heading("Date : "+Listarg[1], level=2)
first_page_firstPart = doc.add_paragraph(contenu_texte)
first_page_firstPart.runs[0].font.size = Pt(10)
first_page_firstPart.runs[0].bold = True
transition = doc.add_paragraph("Comment about observation")
transition.runs[0].font.size = Pt(9)
transition.runs[0].underline = True
first_page_SecondPart = doc.add_paragraph(contenu_texte2)
first_page_SecondPart.runs[0].font.size = Pt(7)
doc.add_page_break()

ToDoFileExec = open(ToDoFile,"r")
linesObs = ToDoFileExec.readlines()
ToDoFileExec.close()

ListTotObs=[]
ListObs=[]
ListImage=[]
nbrPipeline=1
nameSave=""
for lineObs in linesObs:
	lineObs=lineObs.replace('\n','')
	ListObs=lineObs.split()
	ListTotObs+= [ListObs]
	ListImage += [PathJpeg+ListObs[0]+"_"+ListObs[1]+"_"+ListObs[6]+"_"+ListObs[7]+"_singlepulse.jpg"]
	if ListObs[0] == nameSave:
		nbrPipeline+=1
	else: nbrPipeline=1
	nameSave=ListObs[0]
#['SGR1935+2154', '20230506', '332', '/databf/nenufar-pulsar/LT05/2023/05/SGR1935+2154_D20230506T0201_60070_501174_0063_BEAM2_0001.fits', '352', '312', '3', '3', '0.5', '0.5']
# name, data , dm , path ,dm max ,dm min , sigmaF ,sigmaT ,BerrF, BerrT
#images = ['FRB180916_D20221229T1806_59942_500982_0063_BEAM2_0001_singlepulse.jpg', 'FRB180916_D20221230T1801_59943_500983_0063_BEAM2_0001_singlepulse.jpg','FRB180916_D20221231T1801_59944_500984_0063_BEAM2_0001_singlepulse.jpg']  
#print(ListImage)
compteurPipeline=1
compteur = 0 
for image in ListImage:
	ligne = "PIPELINE "+str(compteurPipeline)+" | Obs_Name:  "+ListTotObs[compteur][0]+"  Date:  "+ListTotObs[compteur][1]+"  DM: "+ListTotObs[compteur][2]+" pc/cm3\n DM_range |  DM_max:  "+ListTotObs[compteur][4]+" pc/cm3  DM_min:  "+ListTotObs[compteur][5]+" pc/cm3\nParametre |  SigmaF:  "+ListTotObs[compteur][6]+"  SigmaT:  "+ListTotObs[compteur][7]+"  BerrF:  "+ListTotObs[compteur][8]+"  BerrT:  "+ListTotObs[compteur][9]
	if compteurPipeline == nbrPipeline:
		compteurPipeline = 1
	else:
		compteurPipeline +=1
	paragraph = doc.add_paragraph(ligne)
	paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
	paragraph.runs[0].bold = True
	paragraph.runs[0].font.size = Pt(10)
	im_row = Image.open(image, 'r')
	width, height = im_row.size
	new_size = (int(width*0.4),int(height*0.4))
	resized_image = im_row.resize(new_size)
	crop_box=(0,int(width*0.4)*0.4,int(width*0.4), height*0.4)
	cropped_image = resized_image.crop(crop_box)
	cropped_image.save("/data/pgaspari/WeeklyReport/Result/"+Listarg[2]+"/JPEG/cropped_image"+str(compteur)+".jpg")
	#doc.add_picture(image, width=Cm(15), height=Cm(15))
	doc.add_picture("/data/pgaspari/WeeklyReport/Result/"+Listarg[2]+"/JPEG/cropped_image"+str(compteur)+".jpg", width=Cm(15), height=Cm(15))
	doc.add_page_break()
	compteur+=1
	
doc.save(SavePath)

